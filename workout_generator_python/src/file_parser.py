import os
import sys
import argparse

def parse_txt(filepath):
    # Try different encodings
    encodings = ['utf-8', 'utf-8-sig', 'windows-1256', 'cp1252', 'latin-1']
    for enc in encodings:
        try:
            with open(filepath, 'r', encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise Exception("Could not decode TXT file with standard encodings.")

def parse_docx(filepath):
    import docx
    doc = docx.Document(filepath)
    text_parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            text_parts.append(p.text)
    
    # Also parse tables in word document
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                text_parts.append(" | ".join(row_cells))
                
    return "\n".join(text_parts)

def parse_xlsx(filepath):
    import openpyxl
    wb = openpyxl.load_workbook(filepath, data_only=True)
    text_parts = []
    for sheet in wb.worksheets:
        text_parts.append(f"--- Sheet: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                row_str = " | ".join([str(cell).strip() for cell in row if cell is not None])
                text_parts.append(row_str)
    return "\n".join(text_parts)

def main():
    parser = argparse.ArgumentParser(description="BeastMode File Content Extractor")
    parser.add_argument("--file", type=str, required=True, help="Path to txt, docx, or xlsx file")
    args = parser.parse_args()
    
    filepath = args.file
    if not os.path.exists(filepath):
        print(f"Error: File {filepath} does not exist", file=sys.stderr)
        sys.exit(1)
        
    ext = os.path.splitext(filepath)[1].lower()
    
    try:
        # Set stdout to UTF-8
        try:
            sys.stdout.reconfigure(encoding='utf-8')
        except AttributeError:
            pass
            
        if ext == '.txt':
            content = parse_txt(filepath)
        elif ext == '.docx':
            content = parse_docx(filepath)
        elif ext in ['.xlsx', '.xls']:
            content = parse_xlsx(filepath)
        else:
            print(f"Error: Unsupported file format {ext}", file=sys.stderr)
            sys.exit(1)
            
        print(content)
        
    except Exception as e:
        print(f"Error: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
